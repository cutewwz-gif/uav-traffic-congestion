# -*- coding: utf-8 -*-
"""一键生成完善后的期刊模板 Word 论文。"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

PAPER_DIR = Path(__file__).resolve().parent
TEMPLATE = Path(r"d:\下载\1652432177757386.docx")
BACKUP = TEMPLATE.with_suffix(".backup.docx")
FILLED = PAPER_DIR / "1652432177757386_已填写.docx"
OUT_DOCX = PAPER_DIR / "基于YOLOv12的无人机俯视交通流感知与道路拥堵智能研判系统_完稿.docx"

SOURCE_CANDIDATES = [
    PAPER_DIR / "1652432177757386_已填写_最新.docx",
    PAPER_DIR / "1652432177757386_已填写_含混淆矩阵.docx",
    FILLED,
]

TEMPLATE_JUNK = (
    "式”、“如式所示”等说法",
    "结论需在研究结果与讨论的基础上",
    "正文使用五号宋体，段前首行缩进2字符",
    "文中的公式统一使用Mathtype",
    "参考文献的引用应按照引用顺序",
    "1.使用本模板",
    "2.文章正文篇幅",
    "1 格式要求",
    "2 图和表",
    "遵循“先见文字",
    "图片推荐使用",
    "中文表题使用",
    "示例表如表1所列",
    "公式示例如式",
    "Key word：word1",
    "English title",
    "稿件中文题目",
)


def delete_paragraph(para: Paragraph) -> None:
    para._element.getparent().remove(para._element)


def set_para_text(para: Paragraph, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def remove_template_junk(doc: Document) -> int:
    removed = 0
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if not t:
            continue
        if any(k in t for k in TEMPLATE_JUNK):
            delete_paragraph(para)
            removed += 1
            continue
        if t.startswith("1.1 ") and "参考文献标注" in t:
            delete_paragraph(para)
            removed += 1
        if t.startswith("0 引言") and "双栏" in t:
            delete_paragraph(para)
            removed += 1
    return removed


def patch_content(doc: Document) -> None:
    from generate_paper_pdf import ABSTRACT_EN, SECTIONS

    s52 = next(c for _, sub, c, _ in SECTIONS if sub and sub.startswith("5.2"))
    s53 = next(c for _, sub, c, _ in SECTIONS if sub and sub.startswith("5.3"))
    concl = next(c for sec, _, c, _ in SECTIONS if sec and sec.startswith("6") and c)

    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("Abstract:"):
            set_para_text(para, f"Abstract: {ABSTRACT_EN}")
        elif t.startswith("隔帧推理"):
            set_para_text(para, s53)
        elif t.startswith("自由流") or t.startswith("在自建验证集"):
            set_para_text(para, s52)
        elif t.startswith("本文面向无人机") or t.startswith("结论需在研究"):
            set_para_text(para, concl)


def pick_source() -> Path | None:
    for p in SOURCE_CANDIDATES:
        if p.is_file():
            return p
    return None


def rebuild_from_template() -> Path:
    if not BACKUP.is_file() and not TEMPLATE.is_file():
        raise FileNotFoundError("未找到期刊 Word 模板")
    subprocess.run([sys.executable, str(PAPER_DIR / "fill_template.py")], check=True)
    if not FILLED.is_file():
        raise FileNotFoundError("fill_template 未生成 docx")
    return FILLED


def finalize_doc(doc: Document) -> None:
    from update_docx_figures import (
        cleanup_stale_empty_after_heading,
        fix_table1,
        insert_system_figure,
        remove_drawings,
        remove_old_figure_blocks,
        remove_system_figure_blocks,
    )

    remove_drawings(doc)
    remove_old_figure_blocks(doc)
    remove_system_figure_blocks(doc)
    remove_template_junk(doc)
    patch_content(doc)
    fix_table1(doc)
    insert_system_figure(doc)
    cleanup_stale_empty_after_heading(doc)


def build() -> Path:
    src = pick_source()
    if src is None:
        print("未找到已有 docx，从期刊模板重新生成…")
        src = rebuild_from_template()
    else:
        print(f"基于已有文档完善: {src}")

    doc = Document(str(src))
    n_junk = remove_template_junk(doc)
    finalize_doc(doc)

    out = OUT_DOCX
    try:
        doc.save(str(out))
    except PermissionError:
        out = OUT_DOCX.with_stem(OUT_DOCX.stem + "_v2")
        doc.save(str(out))

    if TEMPLATE.parent.is_dir():
        try:
            shutil.copy2(out, TEMPLATE.parent / out.name)
            print(f"已复制到: {TEMPLATE.parent / out.name}")
        except OSError:
            pass

    print(f"删除模板残留: {n_junk} 处")
    print(f"完稿已生成: {out}")
    return out


if __name__ == "__main__":
    build()
