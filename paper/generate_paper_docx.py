# -*- coding: utf-8 -*-
"""生成 Word 版论文（与 PDF 同内容、同体例）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from generate_paper_pdf import (
    ABSTRACT_CN,
    ABSTRACT_EN,
    KEYWORDS_CN,
    KEYWORDS_EN,
    META_LINE,
    PAPER_TITLE,
    PAPER_TITLE_EN,
    REFERENCES,
    SECTIONS,
)

ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "基于YOLOv12的无人机俯视交通流感知与道路拥堵智能研判系统.docx"

FONT_SONG = "宋体"
FONT_HEI = "黑体"


def _set_run(run, font: str, size: float, bold: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def _para(
    doc: Document,
    text: str,
    *,
    font: str = FONT_SONG,
    size: float = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent: bool = True,
    space_before: float = 0,
    space_after: float = 0,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    _set_run(run, font, size, bold)


def _label_para(doc: Document, label: str, content: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    r1 = p.add_run(label)
    _set_run(r1, FONT_HEI, 11, bold=True)
    r2 = p.add_run(content)
    _set_run(r2, FONT_SONG, 11)


def _add_table(doc: Document, raw: str) -> None:
    rows = [line.split("|") for line in raw.strip().split("\n")]
    rows = [[c.strip() for c in row] for row in rows]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_run(run, FONT_SONG, 10.5, bold=(i == 0))


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    # 标题
    _para(doc, PAPER_TITLE, font=FONT_HEI, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=6)
    _para(doc, PAPER_TITLE_EN, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=8)
    _para(doc, "（作者姓名）1", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    _para(doc, "（1  单位名称，城市  邮编）", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=12)

    # 摘要
    _label_para(doc, "摘  要：", ABSTRACT_CN)
    _label_para(doc, "关键词：", KEYWORDS_CN)
    _para(doc, META_LINE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=8)
    _label_para(doc, "Abstract: ", ABSTRACT_EN)
    _label_para(doc, "Keywords: ", KEYWORDS_EN)

    # 正文
    for sec_title, sub_title, content, formula in SECTIONS:
        if sec_title:
            _para(doc, sec_title, font=FONT_HEI, size=14, indent=False, space_before=10, space_after=6)
        if sub_title:
            _para(doc, sub_title, font=FONT_HEI, size=12, indent=False, space_before=6, space_after=4)
        if content:
            if content.startswith("表1  "):
                _para(doc, content, font=FONT_HEI, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            elif "|" in content and "A畅通" in content:
                _add_table(doc, content)
            else:
                _para(doc, content)
        if formula:
            _para(doc, formula, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    # 参考文献
    _para(doc, "参考文献", font=FONT_HEI, size=14, indent=False, space_before=10, space_after=6)
    for ref in REFERENCES:
        _para(doc, ref, size=10.5, indent=False)

    doc.save(str(OUT_DOCX))
    print(f"已生成: {OUT_DOCX}")


if __name__ == "__main__":
    build_docx()
