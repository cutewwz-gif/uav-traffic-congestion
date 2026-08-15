# -*- coding: utf-8 -*-
"""在期刊 Word 模板上填入夏令营论文内容。"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_paper_pdf import (
    ABSTRACT_CN,
    ABSTRACT_EN,
    KEYWORDS_CN,
    KEYWORDS_EN,
    PAPER_TITLE,
    PAPER_TITLE_EN,
    REFERENCES,
    SECTIONS,
)

TEMPLATE = Path(r"d:\下载\1652432177757386.docx")
BACKUP = TEMPLATE.with_suffix(".backup.docx")
OUT_COPY = Path(__file__).resolve().parent / "1652432177757386_已填写.docx"

# 正文段落（不含表1 pipe 行，表单独处理）
BODY_BLOCKS: list[tuple[str, str | None, str | None, str | None]] = [
    b for b in SECTIONS if not (b[2] and "|" in b[2] and "A畅通" in b[2])
]


def is_red_instruction(paragraph) -> bool:
    for run in paragraph.runs:
        if run.font.color and run.font.color.rgb and str(run.font.color.rgb) == "FF0000":
            return True
    return False


def is_deletable_example_body(text: str) -> bool:
    t = text.strip()
    keys = (
        "正文使用五号宋体，段前首行缩进2字符。以下",
        "正文使用五号宋体，段前首行缩进2字符。正文中英文缩写",
        "如涉及人名，需姓在前",
        "文中的公式统一使用Mathtype",
        "参考文献的引用应按照引用顺序",
        "微波部件表面损坏",
        "严格耦合波分析",
        "式”、“如式所示”等说法",
        "符号及计量单位使",
        "结论需在研究结果与讨论的基础上",
        "Key word：word1",
    )
    return any(k in t for k in keys)


def is_deletable_header(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    keys = (
        "论文格式参考模板",
        "模板使用说明",
        "提交论文时",
        "论文被收录后",
        "稿件中文题目",
        "作者简介",
        "通信作者",
        "E-mail",
        "English title",
        "期刊著录",
        "专著著录",
        "译著著录",
        "会议论文集著录",
        "学位论文著录",
        "专利文献著录",
        "技术标准著录",
        "电子文献著录",
        "示例表如表",
        "图1 中文大图题",
        "Fig. 1",
        "表1 中文表题",
        "Tab.1",
        "（a）中文图",
        "公式示例如式",
        "遵循“先见文字",
        "图片推荐使用",
        "中文表题使用",
        "（文章如有致谢",
        "致谢正文",
        "附录（如有",
        "张三（19xx",
        "李某四（19xx",
    )
    return any(k in t for k in keys)


def clear_para(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clone_para_style(doc: Document, after_para, text: str, style_para) -> None:
    """在 after_para 后插入与 style_para 同格式的段落。"""
    new_p = deepcopy(style_para._element)
    after_para._element.addnext(new_p)
    # 重新定位：从 doc.paragraphs 找最后一个匹配较麻烦，直接操作 xml
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, style_para._parent)
    clear_para(para, text)
    return para


def add_body_paragraph(doc: Document, ref_para, text: str, *, heading: str | None = None) -> None:
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(ref_para._element)
    ref_para._element.addnext(new_p)
    para = Paragraph(new_p, ref_para._parent)
    if heading:
        clear_para(para, heading)
        for run in para.runs:
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            if "0 " in heading or heading[0].isdigit() and "." not in heading[:3]:
                run.font.size = Pt(15)  # 三号≈15pt
            elif heading.count(".") == 1:
                run.font.size = Pt(14)  # 四号
            else:
                run.font.size = Pt(10.5)
        # 再插正文
        new_p2 = deepcopy(ref_para._element)
        para._element.addnext(new_p2)
        para2 = Paragraph(new_p2, ref_para._parent)
        clear_para(para2, text)
        return

    clear_para(para, text)
    para.paragraph_format.first_line_indent = ref_para.paragraph_format.first_line_indent


def fill_table1(doc: Document) -> None:
    """用场景判别表替换模板示例表，并移到正文引用处。"""
    if not doc.tables:
        return
    table = doc.tables[0]
    data = [
        ["场景", "速度特征", "波动特征", "Ct区间", "标签"],
        ["A畅通", "Vavg≥Vhigh", "σdist低", "<60", "畅通"],
        ["B有序等待", "低速", "σdist≤σlow", "20—40", "有序等待"],
        ["C走走停停", "低速", "σdist≥σhigh", "≥60", "拥堵"],
        ["E路口等灯", "静动混合", "静组波动小", "≤30", "路口等灯"],
    ]

    for tc in table._tbl.iter_tcs():
        for tag in ("w:vMerge", "w:gridSpan"):
            for el in tc.xpath(f".//{tag}"):
                el.getparent().remove(el)

    while len(table.rows) < len(data):
        table.add_row()
    while len(table.rows) > len(data):
        delete_row = table.rows[-1]._tr
        delete_row.getparent().remove(delete_row)
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = val
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("表1给出了"):
            anchor = p
            break
    if anchor is None:
        return

    tbl_el = table._tbl
    if tbl_el.getparent() is not None:
        tbl_el.getparent().remove(tbl_el)

    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    def _insert_after(ref: Paragraph) -> Paragraph:
        new_p = OxmlElement("w:p")
        ref._element.addnext(new_p)
        return Paragraph(new_p, ref._parent)

    cap_cn = _insert_after(anchor)
    cap_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_cn.add_run("表1  典型交通场景判别条件")
    cap_en = _insert_after(cap_cn)
    cap_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_en.add_run("Tab.1 Typical traffic scenario criteria")
    cap_en._element.addnext(tbl_el)


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)

    if not BACKUP.exists():
        shutil.copy2(TEMPLATE, BACKUP)

    src = BACKUP if BACKUP.exists() else TEMPLATE
    doc = Document(str(src))

    # 1) 删除红色提示段、模板说明段、示例段（倒序删除避免重复引用）
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if is_red_instruction(para) or is_deletable_header(t):
            delete_paragraph(para)
            continue
        if t.startswith("1.使用本模板") or t.startswith("2.文章正文篇幅"):
            delete_paragraph(para)
            continue
        if "DOI" in t and "10.3969" in t:
            delete_paragraph(para)
            continue
        if "https://ztflhxhma.com" in t:
            delete_paragraph(para)

    # 删除示例图段落与模板正文示例
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if t.startswith("（a）") or t.startswith("（b）"):
            delete_paragraph(para)
            continue
        if is_deletable_example_body(t):
            delete_paragraph(para)
    while len(doc.tables) > 1:
        tbl = doc.tables[-1]._tbl
        tbl.getparent().remove(tbl)

    paras = doc.paragraphs

    # 3) 定位并替换前置部分（按剩余文本模糊匹配）
    def find_startswith(prefix: str):
        for p in paras:
            if p.text.strip().startswith(prefix):
                return p
        return None

    title_p = paras[0] if paras else None
    if title_p:
        clear_para(title_p, PAPER_TITLE)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in title_p.runs:
            r.bold = True
            r.font.name = "黑体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            r.font.size = Pt(22)

    # 找作者行（含逗号、较短）
    for p in paras[1:8]:
        t = p.text.strip()
        if "，" in t and len(t) < 40 and "摘要" not in t and "单位" not in t:
            clear_para(p, "（作者姓名）1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    for p in paras:
        t = p.text.strip()
        if t.startswith("（1.") or t.startswith("1.中国") or "710000" in t:
            clear_para(p, "（1. 单位名称，城市  邮编）")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if t.startswith("2.上海") or t.startswith("2."):
            if "学院" in t or "University" in t:
                delete_paragraph(p)

    abs_p = find_startswith("摘")
    if abs_p:
        clear_para(abs_p, f"摘  要：{ABSTRACT_CN}")

    kw_p = None
    for p in paras:
        if p.text.strip().startswith("关键词") or p.text.strip().startswith("关键"):
            kw_p = p
            break
    # 关键词放在摘要之后
    if abs_p is not None:
        if kw_p is not None:
            delete_paragraph(kw_p)
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        abs_p._element.addnext(new_p)
        kw_p = Paragraph(new_p, abs_p._parent)
        clear_para(kw_p, f"关键词：{KEYWORDS_CN}")

    cls_p = find_startswith("中图分类号")
    if cls_p:
        clear_para(cls_p, "中图分类号：U491.265    文献标识码：A    成果类型：夏令营项目展示")

    # 英文题目
    en_title_p = None
    for p in paras:
        t = p.text.strip()
        if t == "English title" or (
            t.replace("-", "").replace(" ", "").isascii()
            and len(t) > 25
            and "Abstract" not in t
            and "University" not in t
            and "Author" not in t
            and "China" not in t
            and not t.startswith("(")
            and "Key" not in t
        ):
            if en_title_p is None or t == "English title":
                en_title_p = p
    if en_title_p:
        clear_para(en_title_p, PAPER_TITLE_EN)
        en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in en_title_p.runs:
            r.bold = True
            r.font.size = Pt(22)
    elif cls_p is not None:
        en_title_p = cls_p.insert_paragraph_before(PAPER_TITLE_EN)
        en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in paras:
        if p.text.strip().startswith("ZHANG"):
            clear_para(p, "Author Name1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.text.strip().startswith("(1.China") or p.text.strip().startswith("(1."):
            clear_para(p, "(1. University Name, City Postcode, China)")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "School of Mechatronic" in p.text:
            delete_paragraph(p)

    en_abs = find_startswith("Abstract")
    if en_abs:
        clear_para(en_abs, f"Abstract: {ABSTRACT_EN}")

    en_kw = find_startswith("Key word") or find_startswith("Keywords")
    if en_kw:
        clear_para(en_kw, f"Key words: {KEYWORDS_EN}")
    elif en_abs is not None:
        en_kw_p = doc.add_paragraph(f"Key words: {KEYWORDS_EN}")
        en_abs._element.addnext(en_kw_p._element)

    # 4) 删除模板正文示例，从「0 引言/格式要求」到「3 结论」之间
    start_body = end_body = concl_p = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if start_body is None and (
            t.startswith("0 ") or "格式要求" in t or (t.startswith("0") and "引言" in t)
        ):
            start_body = p
        if t.startswith("3 结论") or t.startswith("3  结论") or (t.startswith("3 ") and "结论" in t):
            end_body = p
            concl_p = p
            break

    body_style = start_body
    if start_body and end_body and start_body is not end_body:
        deleting = False
        for para in list(doc.paragraphs):
            if para is start_body:
                deleting = True
                continue
            if para is end_body:
                deleting = False
                continue
            if deleting:
                delete_paragraph(para)

    # 5) 在结论段前插入正文
    if concl_p is None:
        for p in doc.paragraphs:
            if "结论" in p.text.strip()[:6]:
                concl_p = p
                break

    if concl_p is None:
        concl_p = doc.add_paragraph("6  结  论")

    insert_before = concl_p
    ref_para = body_style
    if ref_para is None:
        for p in doc.paragraphs:
            if p.text.strip().startswith("正文使用"):
                ref_para = p
                break
    if ref_para is None:
        ref_para = doc.paragraphs[min(5, len(doc.paragraphs) - 1)]

    # 倒序插入以保持顺序
    blocks: list[tuple[str, str]] = []
    for sec, sub, content, formula in BODY_BLOCKS:
        if sec:
            blocks.append(("h1", sec))
        if sub:
            blocks.append(("h2", sub))
        if content and not content.startswith("表1  "):
            blocks.append(("body", content))
        if formula:
            blocks.append(("formula", formula))

    for kind, text in blocks:
        new_p = deepcopy(ref_para._element)
        insert_before._element.addprevious(new_p)
        from docx.text.paragraph import Paragraph

        para = Paragraph(new_p, ref_para._parent)
        clear_para(para, text)
        if kind == "h1":
            for r in para.runs:
                r.bold = True
                r.font.name = "黑体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                r.font.size = Pt(15)
        elif kind == "h2":
            for r in para.runs:
                r.bold = True
                r.font.name = "黑体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                r.font.size = Pt(14)
        elif kind == "formula":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 删除残留的「0 引言/格式要求」占位段
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("0 引言") and "双栏" in t:
            delete_paragraph(para)
        if t.startswith("1 格式要求"):
            delete_paragraph(para)
        if t.startswith("1.1 ") and "参考文献标注" in t:
            delete_paragraph(para)
        if t.startswith("1.1.1"):
            delete_paragraph(para)
        if t.startswith("2 图和表"):
            delete_paragraph(para)

    # 6) 结论（只保留一节）
    concl_body = ""
    for _, _, c, _ in SECTIONS:
        if c and c.startswith("本文面向无人机垂直俯视"):
            concl_body = c
            break

    seen_concl = False
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if t.startswith("3 结论") or t.startswith("6  结  论") or t == "6  结  论":
            if seen_concl:
                delete_paragraph(para)
                continue
            seen_concl = True
            clear_para(para, "6  结  论")
            for r in para.runs:
                r.bold = True
            continue
        if concl_body and (
            t.startswith("结论需在研究")
            or (t.startswith("本文面向无人机垂直俯视") and seen_concl)
        ):
            clear_para(para, concl_body)
            concl_body = ""  # only once
        if t.startswith("对比时间序列预测") and seen_concl:
            delete_paragraph(para)
        if t.startswith("结论需在研究结果"):
            delete_paragraph(para)

    # 7) 参考文献
    ref_header = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("参考文献"):
            ref_header = p
            break
    if ref_header is None:
        ref_header = doc.add_paragraph("参考文献")
        for r in ref_header.runs:
            r.bold = True
    else:
        clear_para(ref_header, "参考文献")
        for r in ref_header.runs:
            r.bold = True

    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if t.startswith("[1]") and ("张鹏飞" in t or "铝合金" in t or "时柏营" not in t and "张溪" not in t):
            if "张溪" not in t and "时柏营" not in t:
                delete_paragraph(para)
        if t.startswith("[2]") and "李长河" in t:
            delete_paragraph(para)
        if t.startswith("[3]") and "尼葛洛庞" in t:
            delete_paragraph(para)
        if t.startswith("[4]") and "ROMIER" in t:
            delete_paragraph(para)
        if t.startswith("[5]") and "黄桂平" in t:
            delete_paragraph(para)
        if t.startswith("[6]") and "Uher" in t:
            delete_paragraph(para)
        if t.startswith("[7]") and "质量监督" in t:
            delete_paragraph(para)
        if "中国互联网发展" in t:
            delete_paragraph(para)

    # 删除已有项目参考文献避免重复
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if t.startswith("[1] 张溪") or t.startswith("[2] 时柏营"):
            delete_paragraph(para)

    anchor = ref_header
    for ref in REFERENCES:
        new_p = deepcopy(ref_para._element)
        anchor._element.addnext(new_p)
        from docx.text.paragraph import Paragraph

        para = Paragraph(new_p, ref_para._parent)
        clear_para(para, ref)
        anchor = para

    # 8) 表1
    fill_table1(doc)
    for p in doc.paragraphs:
        if p.text.strip() == "表1 中文表题":
            clear_para(p, "表1  典型交通场景判别条件")
        if p.text.strip().startswith("Tab.1"):
            clear_para(p, "Tab.1 Typical traffic scenario criteria")

    doc.save(str(OUT_COPY))
    try:
        doc.save(str(TEMPLATE))
        print(f"已写入模板: {TEMPLATE}")
    except PermissionError:
        print(f"模板文件被占用，未覆盖原文件。请关闭 Word 后重新运行。")
    print(f"备份文件: {BACKUP}")
    print(f"输出文件: {OUT_COPY}")


if __name__ == "__main__":
    main()
