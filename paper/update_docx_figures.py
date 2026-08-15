# -*- coding: utf-8 -*-
"""更新 Word：替换为系统界面截图，修复表1位置与排版。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

DOCX_CANDIDATES = [
    Path(r"D:\finalpr\paper\1652432177757386_已填写_最新.docx"),
    Path(r"D:\finalpr\paper\1652432177757386_已填写_含混淆矩阵.docx"),
    Path(r"D:\finalpr\paper\1652432177757386_已填写.docx"),
]
DOCX_OUT = Path(r"D:\finalpr\paper\1652432177757386_已填写.docx")
UI_IMG = Path(r"D:\finalpr\paper\figures\system_ui.png")

TABLE_DATA = [
    ["场景", "速度特征", "波动特征", "Ct区间", "标签"],
    ["A畅通", "Vavg≥Vhigh", "σdist低", "<60", "畅通"],
    ["B有序等待", "低速", "σdist≤σlow", "20—40", "有序等待"],
    ["C走走停停", "低速", "σdist≥σhigh", "≥60", "拥堵"],
    ["E路口等灯", "静动混合", "静组波动小", "≤30", "路口等灯"],
]


def pick_source() -> Path:
    for p in DOCX_CANDIDATES:
        if p.is_file():
            return p
    raise FileNotFoundError("未找到已填写的 Word 文档")


def delete_paragraph(para: Paragraph) -> None:
    para._element.getparent().remove(para._element)


def insert_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def set_run_font(run, *, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_caption(
    anchor: Paragraph,
    text: str,
    *,
    hei: bool = True,
    size: float = 9,
    bold: bool = False,
) -> Paragraph:
    para = insert_after(anchor)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, name="黑体" if hei else "Times New Roman", size=size, bold=bold)
    return para


def remove_drawings(doc: Document) -> int:
    removed = 0
    for para in doc.paragraphs:
        for d in para._element.xpath(".//w:drawing"):
            d.getparent().remove(d)
            removed += 1
    return removed


def remove_old_figure_blocks(doc: Document) -> None:
    drop_keys = (
        "（a）YOLOv12",
        "（b）YOLOv12",
        "图1  YOLOv12",
        "Fig.1 Confusion matrix",
        "（a）中文图",
        "（b）中文图",
        "图1 中文大图题",
        "Fig. 1 English title",
    )
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if any(t.startswith(k) or t == k for k in drop_keys):
            delete_paragraph(para)


def clear_cell_merges(table) -> None:
    for tc in table._tbl.iter_tcs():
        for tag in ("w:vMerge", "w:gridSpan"):
            for el in tc.xpath(f".//{tag}"):
                el.getparent().remove(el)


def set_cell_text(cell, text: str, *, header: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, name="黑体" if header else "宋体", size=9 if header else 7.5, bold=header)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def apply_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(borders)


def remove_duplicate_table_captions(doc: Document) -> None:
    anchor = find_table1_anchor(doc)
    if anchor is None:
        return
    el = anchor._element
    passed_table = False
    to_delete: list[Paragraph] = []
    while True:
        el = el.getnext()
        if el is None:
            break
        if el.tag.endswith("tbl"):
            passed_table = True
            continue
        if not el.tag.endswith("p"):
            continue
        para = Paragraph(el, anchor._parent)
        t = para.text.strip()
        if passed_table and t in (
            "表1  典型交通场景判别条件",
            "Tab.1 Typical traffic scenario criteria",
        ):
            to_delete.append(para)
    for para in reversed(to_delete):
        delete_paragraph(para)


def find_table1_anchor(doc: Document) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("表1给出了"):
            return p
    return None


def table_follows_captions(anchor: Paragraph) -> bool:
    el = anchor._element
    for _ in range(4):
        el = el.getnext()
        if el is None:
            return False
        if el.tag.endswith("tbl"):
            return True
    return False


def fix_table1(doc: Document) -> None:
    if not doc.tables:
        return
    table = doc.tables[0]
    clear_cell_merges(table)

    while len(table.rows) < len(TABLE_DATA):
        table.add_row()
    while len(table.rows) > len(TABLE_DATA):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    for i, row_data in enumerate(TABLE_DATA):
        for j, val in enumerate(row_data):
            set_cell_text(table.rows[i].cells[j], val, header=(i == 0))

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(table)
    width = Cm(16.0)
    for col in table.columns:
        col.width = int(width / len(table.columns))

    anchor = find_table1_anchor(doc)
    if anchor is None:
        return

    tbl_el = table._tbl
    if table_follows_captions(anchor):
        remove_duplicate_table_captions(doc)
        return

    if tbl_el.getparent() is not None:
        tbl_el.getparent().remove(tbl_el)

    cap_cn = add_caption(anchor, "表1  典型交通场景判别条件", size=9)
    cap_en = add_caption(cap_cn, "Tab.1 Typical traffic scenario criteria", hei=False, size=9)
    cap_en._element.addnext(tbl_el)
    remove_duplicate_table_captions(doc)


def remove_system_figure_blocks(doc: Document) -> None:
    drop_keys = (
        "图1  道路拥堵实时分析系统",
        "Fig.1 Web interface of road congestion",
    )
    for para in reversed(list(doc.paragraphs)):
        t = para.text.strip()
        if any(t.startswith(k) or t == k for k in drop_keys):
            delete_paragraph(para)
        elif para._element.xpath(".//w:drawing") and not t:
            delete_paragraph(para)


def cleanup_stale_empty_after_heading(doc: Document) -> None:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("4.2") and "展示界面" in para.text:
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if (
                nxt is not None
                and not nxt.text.strip()
                and not nxt._element.xpath(".//w:drawing")
            ):
                delete_paragraph(nxt)
            break


def insert_system_figure(doc: Document) -> None:
    if not UI_IMG.is_file():
        raise FileNotFoundError(UI_IMG)

    anchor = None
    for p in doc.paragraphs:
        if "前端布局为左侧控制" in p.text:
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith("4.2") and "展示界面" in t:
                anchor = p
                break
    if anchor is None:
        anchor = doc.paragraphs[min(40, len(doc.paragraphs) - 1)]

    pic = insert_after(anchor)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(UI_IMG), width=Cm(15.5))

    cap1 = add_caption(pic, "图1  道路拥堵实时分析系统 Web 展示界面", size=10.5)
    add_caption(
        cap1,
        "Fig.1 Web interface of road congestion real-time analysis system",
        hei=False,
        size=9,
    )


def save_doc(doc: Document) -> None:
    try:
        doc.save(str(DOCX_OUT))
        print(f"已保存: {DOCX_OUT}")
    except PermissionError:
        alt = DOCX_OUT.with_stem(DOCX_OUT.stem + "_最新")
        doc.save(str(alt))
        print(f"原文件被占用，已另存为: {alt}")


def main() -> None:
    src = pick_source()
    doc = Document(str(src))

    n_img = remove_drawings(doc)
    remove_old_figure_blocks(doc)
    remove_system_figure_blocks(doc)
    fix_table1(doc)
    insert_system_figure(doc)
    cleanup_stale_empty_after_heading(doc)
    save_doc(doc)

    print(f"来源: {src}")
    print(f"清除旧图: {n_img} 处；已插入系统界面图并修复表1排版。")


if __name__ == "__main__":
    main()
